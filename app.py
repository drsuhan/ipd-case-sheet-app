# IPD Case Sheet Generator Web Interface for Dr. Doodley Pet Hospital

import datetime
import streamlit as st
from docx import Document
from io import BytesIO

def generate_ipd_case_sheet_docx(data):
    doc = Document()
    doc.add_heading('Dr. Doodley Pet Hospital, Jayanagar Branch', 0)
    doc.add_heading('In-Patient Department (IPD) Case Sheet', level=1)
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%d-%m-%Y')}")
    doc.add_paragraph("Doctor-in-Charge: " + ', '.join(data['doctor_in_charge']))
    doc.add_paragraph("\n")

    def add_section(title, content_dict):
        doc.add_heading(title, level=2)
        for key, value in content_dict.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.add_paragraph("")

    add_section("1. Patient Details", data['patient_details'])
    add_section("2. Clinical Presentation", data['clinical_presentation'])
    add_section("3. Vitals", data['vitals'])
    add_section("4. Diagnostic Tests Done", data['diagnostic_tests'])

    doc.add_heading("5. Provisional Diagnosis", level=2)
    doc.add_paragraph(data['provisional_diagnosis'])
    doc.add_heading("6. Differential Diagnosis", level=2)
    doc.add_paragraph(data['differential_diagnosis'])
    doc.add_heading("7. Final Diagnosis", level=2)
    doc.add_paragraph(data['final_diagnosis'])

    doc.add_heading("8. Treatment Plan", level=2)
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Drug/Fluid Name'
    hdr_cells[1].text = 'Dosage (mg/ml)'
    hdr_cells[2].text = 'Route'
    hdr_cells[3].text = 'Frequency'
    hdr_cells[4].text = 'Duration'
    for item in data['treatment_plan']:
        row_cells = table.add_row().cells
        row_cells[0].text = item['name']
        row_cells[1].text = item['dosage']
        row_cells[2].text = item['route']
        row_cells[3].text = item['frequency']
        row_cells[4].text = item['duration']

    doc.add_heading("9. Special Procedures / Surgery Done", level=2)
    doc.add_paragraph(data['special_procedures'])

    doc.add_heading("10. Follow-up Instructions", level=2)
    doc.add_paragraph(data['follow_up'])

    doc.add_heading("11. Discharge Summary", level=2)
    doc.add_paragraph(data['discharge_summary'])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.title("Dr. Doodley Pet Hospital - IPD Case Sheet Generator")

with st.form("ipd_form"):
    st.header("1. Patient Details")
    patient_details = {
        "Pet name": st.text_input("Pet name"),
        "Age": st.text_input("Age"),
        "Species & breed": st.text_input("Species & breed"),
        "Sex": st.selectbox("Sex", ["Male", "Female"]),
        "Body weight": st.text_input("Body weight"),
        "Owner name & contact": st.text_input("Owner name & contact")
    }

    st.header("2. Clinical Presentation")
    clinical_presentation = {
        "Date of admission": st.date_input("Date of admission"),
        "Chief complaints": st.text_area("Chief complaints"),
        "History": st.text_area("History")
    }

    st.header("3. Vitals")
    vitals = {
        "Temperature": st.text_input("Temperature"),
        "Heart rate": st.text_input("Heart rate"),
        "Respiratory rate": st.text_input("Respiratory rate"),
        "Mucous membrane color": st.text_input("Mucous membrane color"),
        "CRT": st.text_input("CRT"),
        "Hydration status": st.text_input("Hydration status")
    }

    st.header("4. Diagnostic Tests Done")
    diagnostic_tests = {
        "CBC with interpretation": st.text_area("CBC with interpretation"),
        "LFT": st.text_area("LFT"),
        "KFT": st.text_area("KFT"),
        "X-ray findings": st.text_area("X-ray findings"),
        "USG findings": st.text_area("USG findings"),
        "ECG": st.text_area("ECG"),
        "Others": st.text_area("Others")
    }

    st.header("5-7. Diagnosis")
    provisional_diagnosis = st.text_input("Provisional Diagnosis")
    differential_diagnosis = st.text_input("Differential Diagnosis")
    final_diagnosis = st.text_input("Final Diagnosis")

    st.header("8. Treatment Plan")
    treatment_plan = []
    for i in range(1, 6):
        with st.expander(f"Treatment {i}"):
            name = st.text_input(f"Drug/Fluid Name {i}", key=f"name_{i}")
            dosage = st.text_input(f"Dosage (mg/ml) {i}", key=f"dosage_{i}")
            route = st.text_input(f"Route {i}", key=f"route_{i}")
            frequency = st.text_input(f"Frequency {i}", key=f"frequency_{i}")
            duration = st.text_input(f"Duration {i}", key=f"duration_{i}")
            if name:
                treatment_plan.append({"name": name, "dosage": dosage, "route": route, "frequency": frequency, "duration": duration})

    st.header("9. Special Procedures / Surgery Done")
    special_procedures = st.text_area("Special Procedures")

    st.header("10. Follow-up Instructions")
    follow_up = st.text_area("Follow-up Instructions")

    st.header("11. Discharge Summary")
    discharge_summary = st.text_area("Discharge Summary")

    doctor_in_charge = st.multiselect("Doctor-in-Charge", ["Dr. Revathi", "Dr. Suhan H M, MVSc (Surgery & Radiology), F.VMAS, PGDAW"])

    submitted = st.form_submit_button("Generate Case Sheet")

if submitted:
    docx_file = generate_ipd_case_sheet_docx({
        "patient_details": patient_details,
        "clinical_presentation": clinical_presentation,
        "vitals": vitals,
        "diagnostic_tests": diagnostic_tests,
        "provisional_diagnosis": provisional_diagnosis,
        "differential_diagnosis": differential_diagnosis,
        "final_diagnosis": final_diagnosis,
        "treatment_plan": treatment_plan,
        "special_procedures": special_procedures,
        "follow_up": follow_up,
        "discharge_summary": discharge_summary,
        "doctor_in_charge": doctor_in_charge
    })

    st.success("Case sheet generated successfully!")

    st.download_button(
        label="Download Case Sheet (Word)",
        data=docx_file,
        file_name="IPD_Case_Sheet.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
